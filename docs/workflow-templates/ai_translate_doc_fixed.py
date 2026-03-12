#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
使用 AI 翻译生物医学文档 - 修复版本
解决文件名过长和格式问题
"""

import json
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# 生物医学翻译术语库 (质量控制/微生物检测领域)
TERMINOLOGY = {
    # 核心术语
    "原液": "Drug Substance",
    "微生物限度检测": "Microbial Limit Test",
    "分析方法验证报告": "Analysis Method Validation Report",
    "项目": "Project",
    "验证": "Validation",
    "报告": "Report",
    
    # 文档结构
    "目的": "Purpose",
    "范围": "Scope",
    "定义": "Definition",
    "职责": "Responsibilities",
    "程序": "Procedure",
    "总结": "Summary",
    "结论": "Conclusion",
    "参考文件": "References",
    "附录": "Appendixes",
    "模板和表格": "Templates and Forms",
    "修订历史": "Revision History",
    
    # 微生物检测术语
    "需氧菌总数": "Total Aerobic Microbial Count (TAMC)",
    "霉菌和酵母菌总数": "Total Combined Molds and Yeasts Count (TYMC)",
    "菌落": "Colony",
    "菌悬液": "Bacterial Suspension",
    "培养基": "Culture Medium",
    "培养": "Incubation",
    "回收率": "Recovery Rate",
    "抑菌性": "Bacteriostasis",
    "无菌生长": "No Growth",
    "有菌生长": "Growth Present",
    "阴性对照": "Negative Control",
    "阳性对照": "Positive Control",
    "供试品": "Test Solution",
    "试验组": "Test Group",
    "对照组": "Control Group",
    
    # 设备术语
    "生物安全柜": "Biological Safety Cabinet",
    "生化培养箱": "Incubator",
    "移液器": "Pipette",
    "过滤支架": "Filter Holder",
    "隔膜泵": "Diaphragm Pump",
    
    # 菌种名称
    "金黄色葡萄球菌": "Staphylococcus aureus",
    "铜绿假单胞菌": "Pseudomonas aeruginosa",
    "枯草芽孢杆菌": "Bacillus subtilis",
    "白色念珠菌": "Candida albicans",
    "黑曲霉": "Aspergillus niger",
    
    # 培养基缩写
    "胰酪大豆胨琼脂培养基": "Tryptone Soy Agar (TSA)",
    "沙氏葡萄糖琼脂培养基": "Sabouraud Dextrose Agar (SDA)",
    "PH7.0 无菌氯化钠 - 蛋白胨缓冲液": "Sterile Sodium Chloride-Peptone Buffer pH7.0",
    
    # 部门和职位
    "质量控制部": "Quality Control Department",
    "质量保证部": "Quality Assurance Department",
    "质量运营": "Quality Operations",
    "检验员": "Inspector",
    "经理": "Manager",
    "质量负责人": "Quality Director",
    "起草": "Author",
    "审核": "Reviewed By",
    "批准": "Approved By",
    
    # 其他术语
    "批号": "Batch No.",
    "有效期": "Expiry Date",
    "传代次数": "Passage Number",
    "代数": "Generation",
    "设备编号": "Equipment No.",
    "校准有效期": "Calibration Expiry Date",
    "验证有效期": "Validation Expiry Date",
    "文件编号": "Document No.",
    "版本号": "Version No.",
    "生效日期": "Effective Date",
    "变更描述": "Change Description",
    "新建": "New",
    "中国药典": "Chinese Pharmacopoeia",
    "美国药典": "United States Pharmacopeia",
    "通则": "General Chapter",
    "验证管理规程": "Validation Management Procedure",
    "分析方法验证管理规程": "Analytical Method Validation Management Procedure",
    "标准操作规程": "Standard Operating Procedure",
}


def translate_text(text):
    """
    使用术语库翻译文本
    """
    if not text or not text.strip():
        return text
    
    translated = text
    
    # 按长度降序排序，优先匹配长的术语
    sorted_terms = sorted(TERMINOLOGY.items(), key=lambda x: len(x[0]), reverse=True)
    
    for zh_term, en_term in sorted_terms:
        if zh_term in translated:
            translated = translated.replace(zh_term, en_term)
    
    return translated


def create_translated_document(extracted_json_path, output_path):
    """
    根据提取的 JSON 内容创建翻译后的 Word 文档
    """
    # 读取提取的内容
    with open(extracted_json_path, 'r', encoding='utf-8') as f:
        content = json.load(f)
    
    # 创建新文档
    doc = Document()
    
    # 设置页面布局为 A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # 统计信息
    total_paragraphs = len(content['paragraphs'])
    total_tables = len(content['tables'])
    translated_paragraphs = 0
    translated_tables = 0
    
    print(f"文档包含 {total_paragraphs} 个段落和 {total_tables} 个表格")
    print("\n正在翻译段落...")
    
    # 翻译段落
    for idx, para_info in enumerate(content['paragraphs'], 1):
        text = para_info['text']
        style = para_info.get('style', 'Normal')
        alignment = para_info.get('alignment')
        font_info = para_info.get('font', {})
        
        # 翻译文本
        translated_text = translate_text(text)
        if translated_text != text:
            translated_paragraphs += 1
        
        # 添加段落
        if style == 'Title':
            para = doc.add_heading('', level=0)
        elif style == 'Heading 1':
            para = doc.add_heading('', level=1)
        elif style == 'level 2':
            para = doc.add_heading('', level=2)
        elif 'toc' in style:
            # 目录样式
            para = doc.add_paragraph()
        else:
            para = doc.add_paragraph()
        
        # 设置文本
        run = para.add_run(translated_text)
        
        # 设置字体
        run.font.name = 'Arial'
        if font_info.get('size'):
            try:
                # 转换字号 (从 EMU 到 pt)
                size_emu = font_info['size']
                if 'Pt' in size_emu or 'pt' in size_emu:
                    size_pt = float(size_emu.replace('Pt', '').replace('pt', ''))
                else:
                    size_pt = float(size_emu) / 100  # 近似转换
                run.font.size = Pt(size_pt)
            except:
                run.font.size = Pt(11)
        else:
            run.font.size = Pt(11)
        
        if font_info.get('bold'):
            run.font.bold = True
        if font_info.get('italic'):
            run.font.italic = True
        
        # 设置对齐方式
        if alignment:
            if 'CENTER' in alignment:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'JUSTIFY' in alignment:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 显示进度 (前 10 个和最后 5 个)
        if idx <= 10 or idx > total_paragraphs - 5:
            print(f"  段落 {idx}: {text[:50]}... -> {translated_text[:50]}...")
    
    print(f"段落翻译完成：{translated_paragraphs}/{total_paragraphs}")
    
    # 翻译表格
    print("\n正在翻译表格...")
    for table_idx, table_info in enumerate(content['tables'], 1):
        rows = len(table_info['rows'])
        if rows > 0:
            cols = len(table_info['rows'][0])
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # 填充表格内容
            table_translated = False
            for row_idx, row_data in enumerate(table_info['rows']):
                for col_idx, cell_data in enumerate(row_data):
                    cell = table.cell(row_idx, col_idx)
                    original_text = cell_data['text']
                    translated_text = translate_text(original_text)
                    
                    if translated_text != original_text:
                        table_translated = True
                    
                    cell.text = translated_text
            
            if table_translated:
                translated_tables += 1
                print(f"  表格 {table_idx}: 已翻译")
    
    print(f"表格翻译完成：{translated_tables}/{total_tables}")
    
    # 添加页脚
    print("\n正在添加页脚...")
    footer_translated = 0
    for footer_info in content['footers']:
        footer_text = footer_info['text']
        translated_footer = translate_text(footer_text)
        if translated_footer != footer_text:
            footer_translated += 1
        
        # 在文档末尾添加页脚信息
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(translated_footer)
        run.font.size = Pt(9)
    
    print(f"页脚翻译完成：{footer_translated} 处")
    
    # 保存文档
    print(f"\n正在保存文档：{output_path}")
    try:
        doc.save(output_path)
        print(f"\n✅ 翻译完成！")
        print(f"  输入文件：{extracted_json_path}")
        print(f"  输出文件：{output_path}")
        print(f"  翻译段落：{translated_paragraphs}/{total_paragraphs}")
        print(f"  翻译表格：{translated_tables}/{total_tables}")
        print(f"  翻译页脚：{footer_translated} 处")
    except Exception as e:
        print(f"\n❌ 保存失败：{e}")
        # 尝试使用短文件名保存
        short_name = "ZW062_DS_Microbial_Limit_Validation_Report_EN.docx"
        print(f"尝试使用短文件名保存：{short_name}")
        doc.save(short_name)
        print(f"✅ 已保存到：{short_name}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法：python ai_translate_doc_fixed.py <extracted_content.json> <output.docx>")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    create_translated_document(input_path, output_path)

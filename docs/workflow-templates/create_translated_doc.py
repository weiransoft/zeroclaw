#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成翻译后的 Word 文档
"""

import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_translated_document(extracted_json_path, output_path):
    """
    根据提取的 JSON 内容创建翻译后的 Word 文档
    """
    # 读取提取的内容
    with open(extracted_json_path, 'r', encoding='utf-8') as f:
        content = json.load(f)
    
    # 创建新文档
    doc = Document()
    
    # 设置页面布局为 A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # 翻译段落
    print("正在创建翻译后的段落...")
    for para_info in content['paragraphs']:
        text = para_info['text']
        style = para_info.get('style', 'Normal')
        alignment = para_info.get('alignment')
        font_info = para_info.get('font', {})
        
        # 添加段落
        if style == 'Title':
            para = doc.add_heading('', level=0)
        elif style == 'Heading 1':
            para = doc.add_heading('', level=1)
        elif style == 'level 2':
            para = doc.add_heading('', level=2)
        elif 'toc' in style:
            # 目录样式
            para = doc.add_paragraph()
        else:
            para = doc.add_paragraph()
        
        # 设置文本 (暂时使用原文，实际翻译需要 AI 参与)
        run = para.add_run(text)
        
        # 设置字体
        if font_info.get('name'):
            run.font.name = 'Arial'
        if font_info.get('size'):
            try:
                size_pt = int(font_info['size'].replace('Pt', '').replace('pt', ''))
                run.font.size = Pt(size_pt)
            except:
                pass
        if font_info.get('bold'):
            run.font.bold = True
        if font_info.get('italic'):
            run.font.italic = True
        
        # 设置对齐方式
        if alignment:
            if 'CENTER' in alignment:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'JUSTIFY' in alignment:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # 翻译表格
    print("正在创建翻译后的表格...")
    for table_info in content['tables']:
        rows = len(table_info['rows'])
        if rows > 0:
            cols = len(table_info['rows'][0])
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # 填充表格内容
            for row_idx, row_data in enumerate(table_info['rows']):
                for col_idx, cell_data in enumerate(row_data):
                    cell = table.cell(row_idx, col_idx)
                    cell.text = cell_data['text']
    
    # 保存文档
    print(f"正在保存文档：{output_path}")
    doc.save(output_path)
    print("✅ 翻译文档创建完成！")


if __name__ == '__main__':
    import sys
    
    input_path = sys.argv[1] if len(sys.argv) > 1 else 'extracted_content.json'
    output_path = sys.argv[2] if len(sys.argv) > 2 else 'translated_document.docx'
    
    create_translated_document(input_path, output_path)

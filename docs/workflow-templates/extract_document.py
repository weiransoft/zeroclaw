#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档内容提取脚本
用于提取 Word 文档的内容以便 AI 翻译
"""

import sys
import json
from docx import Document


def extract_document_content(file_path):
    """
    提取 Word 文档的内容和格式信息
    """
    doc = Document(file_path)
    
    content = {
        "paragraphs": [],
        "tables": [],
        "headers": [],
        "footers": []
    }
    
    # 提取段落
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            paragraph_info = {
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else None,
                "alignment": str(para.alignment) if para.alignment is not None else None
            }
            # 获取第一个run的字体信息
            if para.runs:
                run = para.runs[0]
                paragraph_info["font"] = {
                    "name": run.font.name if run.font.name else None,
                    "size": str(run.font.size) if run.font.size else None,
                    "bold": run.font.bold,
                    "italic": run.font.italic
                }
            content["paragraphs"].append(paragraph_info)
    
    # 提取表格
    for table_idx, table in enumerate(doc.tables):
        table_data = {
            "index": table_idx,
            "rows": []
        }
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for col_idx, cell in enumerate(row.cells):
                cell_content = {
                    "text": cell.text,
                    "paragraphs": []
                }
                for para in cell.paragraphs:
                    if para.text.strip():
                        cell_content["paragraphs"].append({
                            "text": para.text,
                            "style": para.style.name if para.style else None
                        })
                row_data.append(cell_content)
            table_data["rows"].append(row_data)
        content["tables"].append(table_data)
    
    # 提取页眉页脚
    for section_idx, section in enumerate(doc.sections):
        # 页眉
        for para in section.header.paragraphs:
            if para.text.strip():
                content["headers"].append({
                    "section": section_idx,
                    "text": para.text
                })
        
        # 页脚
        for para in section.footer.paragraphs:
            if para.text.strip():
                content["footers"].append({
                    "section": section_idx,
                    "text": para.text
                })
    
    return content


def save_content_to_file(content, output_path):
    """保存提取的内容到 JSON 文件"""
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(content, f, ensure_ascii=False, indent=2)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法：python extract_document.py <input.docx> [output.json]")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else input_path.replace('.docx', '_extracted.json')
    
    print(f"正在提取文档内容：{input_path}")
    content = extract_document_content(input_path)
    
    save_content_to_file(content, output_path)
    print(f"✅ 内容已提取并保存到：{output_path}")
    
    # 打印一些基本信息
    print(f"段落数量：{len(content['paragraphs'])}")
    print(f"表格数量：{len(content['tables'])}")
    print(f"页眉数量：{len(content['headers'])}")
    print(f"页脚数量：{len(content['footers'])}")
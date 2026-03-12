#!/usr/bin/env python3
"""验证翻译后的文档"""

from docx import Document

# 验证翻译后的文档
doc = Document('AMVR-2025-0113 ZW062 Drug Substance - Microbial Limit Test Analysis Method Validation Report clean.docx')

print(f"✅ 翻译文档验证成功")
print(f"段落数：{len(doc.paragraphs)}")
print(f"表格数：{len(doc.tables)}")
print("\n前 5 个段落:")
for i, p in enumerate(doc.paragraphs[:5], 1):
    print(f"  {i}: {p.text[:100]}")

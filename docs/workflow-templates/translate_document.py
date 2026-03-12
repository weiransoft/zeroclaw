#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生物医学文档翻译脚本
使用 AI 翻译 Word 文档，保持格式
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def get_font_name(is_chinese=True):
    """获取字体名称"""
    if is_chinese:
        return '宋体'
    else:
        return 'Arial'


def copy_paragraph_format(src_para, dst_para):
    """复制段落格式"""
    # 复制段落对齐方式
    if src_para.alignment is not None:
        dst_para.alignment = src_para.alignment
    
    # 复制段落格式
    src_fmt = src_para.paragraph_format
    dst_fmt = dst_para.paragraph_format
    
    if src_fmt.left_indent is not None:
        dst_fmt.left_indent = src_fmt.left_indent
    if src_fmt.right_indent is not None:
        dst_fmt.right_indent = src_fmt.right_indent
    if src_fmt.first_line_indent is not None:
        dst_fmt.first_line_indent = src_fmt.first_line_indent
    if src_fmt.line_spacing is not None:
        dst_fmt.line_spacing = src_fmt.line_spacing
    if src_fmt.space_before is not None:
        dst_fmt.space_before = src_fmt.space_before
    if src_fmt.space_after is not None:
        dst_fmt.space_after = src_fmt.space_after


def translate_text_ai(text):
    """
    使用 AI 翻译文本
    这里需要根据实际的 AI 接口进行实现
    """
    # TODO: 实现 AI 翻译调用
    # 目前返回原文本
    return text


def translate_document(input_path, output_path):
    """翻译 Word 文档"""
    print(f"正在读取文档：{input_path}")
    doc = Document(input_path)
    
    # 创建新文档
    new_doc = Document()
    
    # 设置页面布局
    for section in doc.sections:
        new_section = new_doc.add_section()
        new_section.page_width = section.page_width
        new_section.page_height = section.page_height
        new_section.top_margin = section.top_margin
        new_section.bottom_margin = section.bottom_margin
        new_section.left_margin = section.left_margin
        new_section.right_margin = section.right_margin
    
    # 翻译段落
    print("正在翻译段落...")
    for para in doc.paragraphs:
        if para.text.strip():
            # 翻译文本
            translated_text = translate_text_ai(para.text)
            
            # 添加新段落
            new_para = new_doc.add_paragraph()
            
            # 复制格式
            copy_paragraph_format(para, new_para)
            
            # 设置文本
            if len(new_para.runs) > 0:
                new_para.runs[0].text = translated_text
            else:
                new_para.add_run(translated_text)
            
            # 复制字体格式
            if len(para.runs) > 0:
                src_run = para.runs[0]
                dst_run = new_para.runs[0]
                
                if src_run.font.name:
                    dst_run.font.name = get_font_name(False)
                if src_run.font.size:
                    dst_run.font.size = src_run.font.size
                if src_run.font.bold:
                    dst_run.font.bold = src_run.font.bold
                if src_run.font.italic:
                    dst_run.font.italic = src_run.font.italic
    
    # 翻译表格
    print("正在翻译表格...")
    for table in doc.tables:
        # 创建新表格
        rows = len(table.rows)
        cols = len(table.columns)
        new_table = new_doc.add_table(rows=rows, cols=cols)
        
        # 复制表格格式
        new_table.style = table.style
        
        # 翻译单元格
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if cell.text.strip():
                    translated_text = translate_text_ai(cell.text)
                    new_cell = new_table.cell(i, j)
                    
                    # 设置文本
                    if len(new_cell.paragraphs) > 0:
                        para = new_cell.paragraphs[0]
                        if len(para.runs) > 0:
                            para.runs[0].text = translated_text
                        else:
                            para.add_run(translated_text)
                    else:
                        new_cell.add_paragraph(translated_text)
    
    # 保存文档
    print(f"正在保存文档：{output_path}")
    new_doc.save(output_path)
    
    print("✅ 翻译完成！")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法：python translate_document.py <input.docx> <output.docx>")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    translate_document(input_path, output_path)

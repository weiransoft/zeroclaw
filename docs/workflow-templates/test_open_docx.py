#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""测试打开文件"""

from pathlib import Path
from docx import Document

# 测试文件
test_files = [
    'AMVR-2025-0113 ZW062 项目原液 - 微生物限度检测分析方法验证报告 clean.docx',
    'ZW062_DS_Microbial_Limit_Validation_Report_EN.docx'
]

for test_file in test_files:
    p = Path(test_file)
    print(f"文件：{test_file}")
    print(f"  存在：{p.exists()}")
    print(f"  绝对路径：{p.absolute()}")
    
    if p.exists():
        try:
            doc = Document(str(p))
            print(f"  段落数：{len(doc.paragraphs)}")
            print(f"  表格数：{len(doc.tables)}")
            print(f"  ✅ 可以打开")
        except Exception as e:
            print(f"  ❌ 打开失败：{e}")
    print()

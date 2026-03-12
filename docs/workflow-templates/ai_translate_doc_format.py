#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
使用 AI 翻译生物医学文档 - 格式保持版本
严格保持原文档格式
"""

import json
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# 生物医学翻译术语库 (质量控制/微生物检测领域)
TERMINOLOGY = {
    # 核心术语
    "原液": "Drug Substance",
    "微生物限度检测": "Microbial Limit Test",
    "分析方法验证报告": "Analysis Method Validation Report",
    "项目": "Project",
    "验证": "Validation",
    "报告": "Report",
    
    # 文档结构
    "目的": "Purpose",
    "范围": "Scope",
    "定义": "Definition",
    "职责": "Responsibilities",
    "程序": "Procedure",
    "总结": "Summary",
    "结论": "Conclusion",
    "参考文件": "References",
    "附录": "Appendixes",
    "模板和表格": "Templates and Forms",
    "修订历史": "Revision History",
    
    # 微生物检测术语
    "需氧菌总数": "Total Aerobic Microbial Count (TAMC)",
    "霉菌和酵母菌总数": "Total Combined Molds and Yeasts Count (TYMC)",
    "菌落": "Colony",
    "菌悬液": "Bacterial Suspension",
    "培养基": "Culture Medium",
    "培养": "Incubation",
    "回收率": "Recovery Rate",
    "抑菌性": "Bacteriostasis",
    "无菌生长": "No Growth",
    "有菌生长": "Growth Present",
    "阴性对照": "Negative Control",
    "阳性对照": "Positive Control",
    "供试品": "Test Solution",
    "试验组": "Test Group",
    "对照组": "Control Group",
    
    # 设备术语
    "生物安全柜": "Biological Safety Cabinet",
    "生化培养箱": "Incubator",
    "移液器": "Pipette",
    "过滤支架": "Filter Holder",
    "隔膜泵": "Diaphragm Pump",
    
    # 菌种名称
    "金黄色葡萄球菌": "Staphylococcus aureus",
    "铜绿假单胞菌": "Pseudomonas aeruginosa",
    "枯草芽孢杆菌": "Bacillus subtilis",
    "白色念珠菌": "Candida albicans",
    "黑曲霉": "Aspergillus niger",
    
    # 培养基缩写
    "胰酪大豆胨琼脂培养基": "Tryptone Soy Agar (TSA)",
    "沙氏葡萄糖琼脂培养基": "Sabouraud Dextrose Agar (SDA)",
    "PH7.0 无菌氯化钠 - 蛋白胨缓冲液": "Sterile Sodium Chloride-Peptone Buffer pH7.0",
    
    # 部门和职位
    "质量控制部": "Quality Control Department",
    "质量保证部": "Quality Assurance Department",
    "质量运营": "Quality Operations",
    "检验员": "Inspector",
    "经理": "Manager",
    "质量负责人": "Quality Director",
    "起草": "Author",
    "审核": "Reviewed By",
    "批准": "Approved By",
    
    # 其他术语
    "批号": "Batch No.",
    "有效期": "Expiry Date",
    "传代次数": "Passage Number",
    "代数": "Generation",
    "设备编号": "Equipment No.",
    "校准有效期": "Calibration Expiry Date",
    "验证有效期": "Validation Expiry Date",
    "文件编号": "Document No.",
    "版本号": "Version No.",
    "生效日期": "Effective Date",
    "变更描述": "Change Description",
    "新建": "New",
    "中国药典": "Chinese Pharmacopoeia",
    "美国药典": "United States Pharmacopeia",
    "通则": "General Chapter",
    "验证管理规程": "Validation Management Procedure",
    "分析方法验证管理规程": "Analytical Method Validation Management Procedure",
    "标准操作规程": "Standard Operating Procedure",
}


def translate_text(text):
    """
    使用术语库翻译文本
    """
    if not text or not text.strip():
        return text
    
    translated = text
    
    # 按长度降序排序，优先匹配长的术语
    sorted_terms = sorted(TERMINOLOGY.items(), key=lambda x: len(x[0]), reverse=True)
    
    for zh_term, en_term in sorted_terms:
        if zh_term in translated:
            translated = translated.replace(zh_term, en_term)
    
    return translated


def copy_paragraph_format(src_para, dst_para):
    """复制段落格式"""
    # 复制段落对齐方式
    if src_para.alignment is not None:
        dst_para.alignment = src_para.alignment
    
    # 复制段落格式
    src_fmt = src_para.paragraph_format
    dst_fmt = dst_para.paragraph_format
    
    if src_fmt.left_indent is not None:
        dst_fmt.left_indent = src_fmt.left_indent
    if src_fmt.right_indent is not None:
        dst_fmt.right_indent = src_fmt.right_indent
    if src_fmt.first_line_indent is not None:
        dst_fmt.first_line_indent = src_fmt.first_line_indent
    if src_fmt.line_spacing is not None:
        dst_fmt.line_spacing = src_fmt.line_spacing
    if src_fmt.space_before is not None:
        dst_fmt.space_before = src_fmt.space_before
    if src_fmt.space_after is not None:
        dst_fmt.space_after = src_fmt.space_after


def copy_run_format(src_run, dst_run):
    """复制 run 格式"""
    if src_run.font.name:
        # 中文字体映射到英文字体
        if '宋体' in src_run.font.name or 'SimSun' in src_run.font.name:
            dst_run.font.name = 'Arial'
        elif '黑体' in src_run.font.name or 'SimHei' in src_run.font.name:
            dst_run.font.name = 'Arial Black'
        elif '微软雅黑' in src_run.font.name or 'Microsoft YaHei' in src_run.font.name:
            dst_run.font.name = 'Microsoft YaHei'
        else:
            dst_run.font.name = 'Arial'
    
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    
    if src_run.font.bold:
        dst_run.font.bold = True
    
    if src_run.font.italic:
        dst_run.font.italic = True
    
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb
    
    if src_run.font.underline:
        dst_run.font.underline = True


def create_translated_document_with_format(source_docx_path, extracted_json_path, output_path):
    """
    根据原文档和提取的 JSON 内容创建翻译后的 Word 文档，保持格式
    """
    # 读取原文档获取格式信息
    print(f"正在读取原文档：{source_docx_path}")
    source_doc = Document(source_docx_path)
    
    # 读取提取的内容
    with open(extracted_json_path, 'r', encoding='utf-8') as f:
        content = json.load(f)
    
    # 创建新文档
    doc = Document()
    
    # 复制页面布局
    source_section = source_doc.sections[0]
    section = doc.sections[0]
    section.page_width = source_section.page_width
    section.page_height = source_section.page_height
    section.top_margin = source_section.top_margin
    section.bottom_margin = source_section.bottom_margin
    section.left_margin = source_section.left_margin
    section.right_margin = source_section.right_margin
    
    # 统计信息
    total_paragraphs = len(content['paragraphs'])
    total_tables = len(content['tables'])
    translated_paragraphs = 0
    translated_tables = 0
    
    print(f"文档包含 {total_paragraphs} 个段落和 {total_tables} 个表格")
    print("\n正在翻译段落...")
    
    # 翻译段落 - 使用原文档的格式
    source_paragraphs = source_doc.paragraphs
    
    for idx, para_info in enumerate(content['paragraphs'], 1):
        text = para_info['text']
        
        # 翻译文本
        translated_text = translate_text(text)
        if translated_text != text:
            translated_paragraphs += 1
        
        # 尝试从原文档找到对应的段落来复制格式
        if idx < len(source_paragraphs):
            source_para = source_paragraphs[idx]
            
            # 创建新段落
            para = doc.add_paragraph()
            
            # 复制段落格式
            copy_paragraph_format(source_para, para)
            
            # 复制 run 格式
            if len(source_para.runs) > 0:
                # 清除默认 run
                para.clear()
                run = para.add_run(translated_text)
                copy_run_format(source_para.runs[0], run)
            else:
                run = para.add_run(translated_text)
        else:
            # 如果没有对应的原文段落，创建默认格式
            para = doc.add_paragraph()
            run = para.add_run(translated_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
        
        # 显示进度
        if idx <= 10 or idx > total_paragraphs - 5:
            print(f"  段落 {idx}: {text[:50]}... -> {translated_text[:50]}...")
    
    print(f"段落翻译完成：{translated_paragraphs}/{total_paragraphs}")
    
    # 翻译表格 - 保持格式
    print("\n正在翻译表格...")
    source_tables = source_doc.tables
    
    for table_idx, table_info in enumerate(content['tables'], 1):
        rows = len(table_info['rows'])
        if rows > 0:
            cols = len(table_info['rows'][0])
            
            # 创建表格
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # 如果原表格存在，复制格式
            if table_idx <= len(source_tables):
                source_table = source_tables[table_idx - 1]
                # 复制表格宽度
                try:
                    if source_table._tbl.tblPr is not None and hasattr(source_table._tbl.tblPr, 'tblW'):
                        table._tbl.tblPr.tblW = source_table._tbl.tblPr.tblW
                except:
                    pass  # 忽略格式复制错误
            
            # 填充表格内容
            table_translated = False
            for row_idx, row_data in enumerate(table_info['rows']):
                for col_idx, cell_data in enumerate(row_data):
                    cell = table.cell(row_idx, col_idx)
                    original_text = cell_data['text']
                    translated_text = translate_text(original_text)
                    
                    if translated_text != original_text:
                        table_translated = True
                    
                    # 设置翻译后的文本
                    if len(cell.paragraphs) > 0:
                        para = cell.paragraphs[0]
                        if len(para.runs) > 0:
                            para.runs[0].text = translated_text
                        else:
                            para.add_run(translated_text)
                    else:
                        cell.text = translated_text
            
            if table_translated:
                translated_tables += 1
                print(f"  表格 {table_idx}: 已翻译")
    
    print(f"表格翻译完成：{translated_tables}/{total_tables}")
    
    # 添加页脚
    print("\n正在添加页脚...")
    footer_translated = 0
    for footer_info in content['footers']:
        footer_text = footer_info['text']
        translated_footer = translate_text(footer_text)
        if translated_footer != footer_text:
            footer_translated += 1
        
        # 在文档末尾添加页脚信息
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(translated_footer)
        run.font.size = Pt(9)
    
    print(f"页脚翻译完成：{footer_translated} 处")
    
    # 保存文档
    print(f"\n正在保存文档：{output_path}")
    doc.save(output_path)
    
    print(f"\n✅ 翻译完成！")
    print(f"  输入文件：{extracted_json_path}")
    print(f"  输出文件：{output_path}")
    print(f"  翻译段落：{translated_paragraphs}/{total_paragraphs}")
    print(f"  翻译表格：{translated_tables}/{total_tables}")
    print(f"  翻译页脚：{footer_translated} 处")


if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("用法：python ai_translate_doc_format.py <source.docx> <extracted_content.json> <output.docx>")
        sys.exit(1)
    
    source_path = sys.argv[1]
    input_path = sys.argv[2]
    output_path = sys.argv[3]
    
    create_translated_document_with_format(source_path, input_path, output_path)

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
完整翻译生物医学文档 - 彻底翻译所有中文内容
严格保持原文档格式
"""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# 完整术语库 - 质量控制/微生物检测领域
TERMINOLOGY = {
    # 核心术语
    "原液": "Drug Substance",
    "微生物限度检测": "Microbial Limit Test",
    "分析方法验证报告": "Analysis Method Validation Report",
    "项目": "Project",
    "验证": "Validation",
    "报告": "Report",
    "方案": "Protocol",
    "检查": "Test",
    "确认": "Confirm",
    
    # 文档结构
    "目的": "Purpose",
    "范围": "Scope",
    "定义": "Definition",
    "职责": "Responsibilities",
    "程序": "Procedure",
    "总结": "Summary",
    "结论": "Conclusion",
    "参考文件": "References",
    "附录": "Appendixes",
    "模板和表格": "Templates and Forms",
    "修订历史": "Revision History",
    "目录": "Table of Contents",
    
    # 微生物检测术语
    "需氧菌总数": "Total Aerobic Microbial Count (TAMC)",
    "霉菌和酵母菌总数": "Total Combined Molds and Yeasts Count (TYMC)",
    "菌落": "Colony",
    "菌悬液": "Bacterial Suspension",
    "培养基": "Culture Medium",
    "培养": "Incubation",
    "回收率": "Recovery Rate",
    "抑菌性": "Bacteriostasis",
    "无菌生长": "No Growth",
    "有菌生长": "Growth Present",
    "阴性对照": "Negative Control",
    "阳性对照": "Positive Control",
    "供试品": "Test Solution",
    "试验组": "Test Group",
    "对照组": "Control Group",
    "杂菌": "Contamination",
    "本底菌": "Background Bacteria",
    
    # 设备和材料
    "设备": "Equipment",
    "仪器": "Instrument",
    "试剂": "Reagent",
    "耗材": "Consumables",
    "生物安全柜": "Biological Safety Cabinet",
    "生化培养箱": "Incubator",
    "移液器": "Pipette",
    "过滤支架": "Filter Holder",
    "隔膜泵": "Diaphragm Pump",
    "一次性枪头": "Disposable Tip",
    "过滤杯": "Filter Cup",
    "保存杯": "Storage Cup",
    "涂布棒": "Spreader",
    
    # 菌种名称
    "金黄色葡萄球菌": "Staphylococcus aureus",
    "铜绿假单胞菌": "Pseudomonas aeruginosa",
    "枯草芽孢杆菌": "Bacillus subtilis",
    "白色念珠菌": "Candida albicans",
    "黑曲霉": "Aspergillus niger",
    "定量菌株": "Quantitative Strain",
    
    # 培养基
    "胰酪大豆胨琼脂培养基": "Tryptone Soy Agar (TSA)",
    "沙氏葡萄糖琼脂培养基": "Sabouraud Dextrose Agar (SDA)",
    "PH7.0 无菌氯化钠 - 蛋白胨缓冲液": "Sterile Sodium Chloride-Peptone Buffer pH7.0",
    
    # 部门和职位
    "质量控制部": "Quality Control Department",
    "质量保证部": "Quality Assurance Department",
    "质量运营": "Quality Operations",
    "检验员": "Inspector",
    "经理": "Manager",
    "质量负责人": "Quality Director",
    "起草": "Author",
    "审核": "Reviewed By",
    "批准": "Approved By",
    
    # 其他术语
    "批号": "Batch No.",
    "有效期": "Expiry Date",
    "传代次数": "Passage Number",
    "代数": "Generation",
    "设备编号": "Equipment No.",
    "校准有效期": "Calibration Expiry Date",
    "验证有效期": "Validation Expiry Date",
    "文件编号": "Document No.",
    "版本号": "Version No.",
    "生效日期": "Effective Date",
    "变更描述": "Change Description",
    "新建": "New",
    "中国药典": "Chinese Pharmacopoeia",
    "美国药典": "United States Pharmacopeia",
    "通则": "General Chapter",
    "验证管理规程": "Validation Management Procedure",
    "分析方法验证管理规程": "Analytical Method Validation Management Procedure",
    "标准操作规程": "Standard Operating Procedure",
    "生物安全柜标准操作规程": "Biological Safety Cabinet Standard Operating Procedure",
    "生化培养箱标准操作规程": "Incubator Standard Operating Procedure",
    "文件确认表": "Document Confirmation Form",
    "试剂、菌株、耗材确认表": "Reagents, Strains, Consumables Confirmation Form",
    "仪器确认表": "Equipment Confirmation Form",
    "签名确认表": "Signature Confirmation Form",
    "验证记录表": "Validation Record Form",
    "偏差列表": "Deviation List",
    "方案变更表": "Protocol Change Form",
    
    # 公司和产品
    "上海兆维医药科技有限公司": "Shanghai Hongene Biologics Tech Co., Ltd.",
    "Hongene": "Hongene",
    
    # 常用词
    "本": "This",
    "该": "This",
    "进行": "Conduct",
    "适用": "Applicable",
    "生产": "Produced by",
    "描述": "Describe",
    "条件": "Condition",
    "下": "Under",
    "时": "When",
    "以": "To",
    "确保": "Ensure",
    "可以": "Can",
    "稳定": "Stable",
    "地": "",
    "应用": "Apply to",
    "的": "",
    "及": "and",
    "与": "and",
    "或": "or",
    "在": "In",
    "为": "As",
    "均": "All",
    "已": "Already",
    "经": "Have been",
    "过": "",
    "且": "And",
    "无": "No",
    "没有": "No",
    "所有": "All",
    "要求": "Requirement",
    "的": "of",
    "得到": "Met",
    "符合": "Meet",
    "规定": "Requirement",
    "之间": "Between",
    "生长": "Growth",
    "良好": "Well",
    "未": "Not",
    "超过": "Exceed",
    "代": "Generation",
    "内": "Within",
    "使用": "Use",
    "加入": "Add",
    "量": "Volume",
    "的": "",
    "菌": "Bacteria",
    "计算": "Calculate",
    "出": "Out",
    "并": "And",
    "确定": "Determine",
    "所": "",
    "加": "Added",
    "数量": "Quantity",
    "用于": "Used for",
    "测定": "Determine",
    "中": "In",
    "总数": "Total Count",
    "和": "and",
    "确定": "Determine",
    "整个": "Entire",
    "操作": "Operation",
    "是": "Is",
    "无菌": "Sterile",
    "准备": "Prepare",
    "从": "From",
    "对应": "Corresponding",
    "储存": "Storage",
    "取出": "Take out",
    "待测": "To be tested",
    "样品": "Sample",
    "室温": "Room temperature",
    "静置": "Let stand",
    "使其": "Make it",
    "解冻": "Thaw",
    "融化": "Melt",
    "先": "First",
    "用": "Use",
    "润湿": "Moisten",
    "滤膜": "Filter membrane",
    "过滤": "Filter",
    "再": "Then",
    "滤干": "Filter dry",
    "后": "After",
    "用": "Use",
    "镊子": "Tweezers",
    "取出": "Take out",
    "菌面": "Bacterial side",
    "朝上": "Up",
    "贴于": "Attach to",
    "相对应": "Corresponding",
    "上": "On",
    "按": "According to",
    "表中": "In the table",
    "规定": "Specified",
    "对": "To",
    "条件": "Condition",
    "进行": "Conduct",
    "以": "As",
    "得到": "Obtain",
    "的": "",
    "菌落数": "Colony count",
    "为": "As",
    "需氧菌": "Aerobic bacteria",
    "如果": "If",
    "在": "On",
    "上": "",
    "发现": "Find",
    "的": "",
    "真菌": "Fungi",
    "也": "Also",
    "按需氧菌总数计数": "Count as total aerobic microbial count",
    "细菌": "Bacteria",
    "也按真菌总数计数": "Also count as total molds count",
    "按真菌总数计数": "Count as total molds count",
    "本次": "This",
    "中": "",
    "无": "No",
    "变更": "Change",
    "偏差": "Deviation",
    "发生": "Occur",
    "已经": "Have been",
    "按照": "According to",
    "预先": "Pre-",
    "批准": "Approved",
    "验证": "Validated",
    "方案文件编号": "Protocol document No.",
    "所有": "All",
    "要求": "Required",
    "活动": "Activities",
    "成功": "Successfully",
    "执行": "Executed",
    "任何": "Any",
    "接受标准": "Acceptance criteria",
    "完全": "Fully",
    "总结": "Summarized",
    "在": "In",
    "中": ""
}


def translate_text(text):
    """
    彻底翻译文本，将所有中文翻译成英文
    """
    if not text or not text.strip():
        return text
    
    translated = text
    
    # 按长度降序排序，优先匹配长的术语
    sorted_terms = sorted(TERMINOLOGY.items(), key=lambda x: len(x[0]), reverse=True)
    
    # 在替换之前，先在中文字符前后添加空格以避免单词连接
    # 这样可以确保中文术语被正确替换而不是合并
    for zh_term, en_term in sorted_terms:
        if zh_term in translated:
            # 替换时确保英文术语周围有适当的空格
            translated = translated.replace(zh_term, f' {en_term} ')
    
    # 清理多余的空格
    import re
    translated = re.sub(r'\s+', ' ', translated)  # 多个空格变一个
    translated = translated.strip()
    
    # 检查是否还有中文字符
    if re.search(r'[\u4e00-\u9fff]', translated):
        # 如果还有中文，尝试逐词翻译
        remaining_chinese = re.findall(r'[\u4e00-\u9fff]+', translated)
        for zh_word in remaining_chinese:
            # 查找最匹配的术语
            for zh_term, en_term in TERMINOLOGY.items():
                if zh_term == zh_word:
                    translated = translated.replace(zh_word, f' {en_term} ')
                    break
    
    # 再次清理多余的空格
    translated = re.sub(r'\s+', ' ', translated)
    translated = translated.strip()
    
    return translated


def copy_paragraph_format(src_para, dst_para):
    """复制段落格式"""
    if src_para.alignment is not None:
        dst_para.alignment = src_para.alignment
    
    src_fmt = src_para.paragraph_format
    dst_fmt = dst_para.paragraph_format
    
    if src_fmt.left_indent is not None:
        dst_fmt.left_indent = src_fmt.left_indent
    if src_fmt.right_indent is not None:
        dst_fmt.right_indent = src_fmt.right_indent
    if src_fmt.first_line_indent is not None:
        dst_fmt.first_line_indent = src_fmt.first_line_indent
    if src_fmt.line_spacing is not None:
        dst_fmt.line_spacing = src_fmt.line_spacing
    if src_fmt.space_before is not None:
        dst_fmt.space_before = src_fmt.space_before
    if src_fmt.space_after is not None:
        dst_fmt.space_after = src_fmt.space_after


def copy_run_format(src_run, dst_run):
    """复制 run 格式"""
    if src_run.font.name:
        dst_run.font.name = 'Arial'
    
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    
    if src_run.font.bold:
        dst_run.font.bold = True
    
    if src_run.font.italic:
        dst_run.font.italic = True
    
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb
    
    if src_run.font.underline:
        dst_run.font.underline = True


def create_fully_translated_document(source_docx_path, output_path):
    """
    完整翻译文档，严格保持格式
    """
    print(f"正在读取原文档：{source_docx_path}")
    source_doc = Document(source_docx_path)
    
    # 创建新文档
    doc = Document()
    
    # 复制页面布局
    source_section = source_doc.sections[0]
    section = doc.sections[0]
    section.page_width = source_section.page_width
    section.page_height = source_section.page_height
    section.top_margin = source_section.top_margin
    section.bottom_margin = source_section.bottom_margin
    section.left_margin = source_section.left_margin
    section.right_margin = source_section.right_margin
    
    # 统计信息
    total_paragraphs = len(source_doc.paragraphs)
    total_tables = len(source_doc.tables)
    translated_paragraphs = 0
    translated_tables = 0
    
    print(f"原文档包含 {total_paragraphs} 个段落和 {total_tables} 个表格")
    print("\n正在翻译段落...")
    
    # 翻译段落
    for idx, source_para in enumerate(source_doc.paragraphs, 1):
        text = source_para.text
        
        # 翻译文本
        translated_text = translate_text(text)
        if translated_text != text:
            translated_paragraphs += 1
        
        # 创建新段落
        para = doc.add_paragraph()
        
        # 复制段落格式
        copy_paragraph_format(source_para, para)
        
        # 复制 run 格式
        if len(source_para.runs) > 0:
            para.clear()
            run = para.add_run(translated_text)
            copy_run_format(source_para.runs[0], run)
        else:
            run = para.add_run(translated_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
        
        # 显示进度
        if idx <= 15 or idx > total_paragraphs - 5:
            print(f"  段落 {idx}: {text[:60]}... -> {translated_text[:60]}...")
    
    print(f"段落翻译完成：{translated_paragraphs}/{total_paragraphs}")
    
    # 翻译表格
    print("\n正在翻译表格...")
    for table_idx, source_table in enumerate(source_doc.tables, 1):
        rows = len(source_table.rows)
        cols = len(source_table.columns)
        
        # 创建表格
        table = doc.add_table(rows=rows, cols=cols)
        table.style = source_table.style if source_table.style else 'Table Grid'
        
        # 填充表格内容
        table_translated = False
        for row_idx, source_row in enumerate(source_table.rows):
            for col_idx, source_cell in enumerate(source_row.cells):
                original_text = source_cell.text
                translated_text = translate_text(original_text)
                
                if translated_text != original_text:
                    table_translated = True
                
                # 设置翻译后的文本
                cell = table.cell(row_idx, col_idx)
                if len(cell.paragraphs) > 0:
                    para = cell.paragraphs[0]
                    if len(para.runs) > 0:
                        para.runs[0].text = translated_text
                    else:
                        para.add_run(translated_text)
                else:
                    cell.text = translated_text
        
        if table_translated:
            translated_tables += 1
            print(f"  表格 {table_idx}: 已翻译")
    
    print(f"表格翻译完成：{translated_tables}/{total_tables}")
    
    # 翻译页眉页脚
    print("\n正在翻译页眉页脚...")
    footer_translated = 0
    for section in doc.sections:
        # 页眉
        for para in section.header.paragraphs:
            if para.text.strip():
                original_text = para.text
                translated_text = translate_text(original_text)
                if translated_text != original_text:
                    footer_translated += 1
                    para.text = translated_text
        
        # 页脚
        for para in section.footer.paragraphs:
            if para.text.strip():
                original_text = para.text
                translated_text = translate_text(original_text)
                if translated_text != original_text:
                    footer_translated += 1
                    para.text = translated_text
    
    print(f"页眉页脚翻译完成：{footer_translated} 处")
    
    # 保存文档
    print(f"\n正在保存文档：{output_path}")
    doc.save(output_path)
    
    print(f"\n✅ 翻译完成！")
    print(f"  输入文件：{source_docx_path}")
    print(f"  输出文件：{output_path}")
    print(f"  翻译段落：{translated_paragraphs}/{total_paragraphs}")
    print(f"  翻译表格：{translated_tables}/{total_tables}")
    print(f"  翻译页眉页脚：{footer_translated} 处")
    
    # 验证是否还有中文
    print("\n验证翻译质量...")
    remaining_chinese = 0
    for idx, para in enumerate(doc.paragraphs, 1):
        if re.search(r'[\u4e00-\u9fff]', para.text):
            remaining_chinese += 1
            if remaining_chinese <= 5:
                print(f"  警告：段落 {idx} 仍有中文：{para.text[:50]}...")
    
    if remaining_chinese > 5:
        print(f"  ... 还有 {remaining_chinese - 5} 个段落包含中文")
    elif remaining_chinese > 0:
        print(f"  共有 {remaining_chinese} 个段落包含未翻译中文")
    else:
        print(f"  ✅ 所有段落都已完全翻译！")


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 3:
        print("用法：python translate_complete.py <source.docx> <output.docx>")
        sys.exit(1)
    
    source_path = sys.argv[1]
    output_path = sys.argv[2]
    
    create_fully_translated_document(source_path, output_path)
